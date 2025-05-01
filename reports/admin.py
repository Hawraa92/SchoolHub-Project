import datetime
import json
import csv
import logging
from io import BytesIO
from django.contrib import admin
from django.http import HttpResponse
from django.urls import reverse
from django.utils.html import format_html, mark_safe
from django.shortcuts import redirect
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from .models import Report, ReportsDashboard

logger = logging.getLogger(__name__)

def export_reports_pdf(modeladmin, request, queryset):
    """
    Admin action to export selected reports as PDF.
    """
    try:
        response = HttpResponse(content_type='application/pdf')
        filename = f'reports_{datetime.datetime.now().strftime("%Y%m%d")}.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        pdf = canvas.Canvas(response, pagesize=letter)
        pdf.setTitle("Reports Information")
        y = 750

        for report in queryset:
            if y < 100:
                pdf.showPage()
                y = 750
            pdf.setFont("Helvetica-Bold", 12)
            pdf.drawString(100, y, f"Report for {report.student.full_name} - {report.get_category_display()}")
            y -= 20
            pdf.setFont("Helvetica", 10)
            pdf.drawString(100, y, f"Report Type: {report.report_type} | Status: {report.status}")
            y -= 15
            pdf.drawString(100, y, f"Generated At: {report.generated_at}")
            y -= 25

            # Attempt to print report data (assuming it is a dict)
            if isinstance(report.data, dict):
                for key, value in report.data.items():
                    pdf.drawString(120, y, f"{key}: {value}")
                    y -= 15
                    if y < 100:
                        pdf.showPage()
                        y = 750
            else:
                pdf.drawString(120, y, f"Data: {report.data}")
                y -= 15
            y -= 20

        pdf.save()
        return response
    except Exception as e:
        logger.exception("Error exporting reports to PDF: %s", e)
        return HttpResponse("An error occurred while generating the PDF.", status=500)

def export_reports_csv(modeladmin, request, queryset):
    """
    Admin action to export selected reports as CSV.
    """
    try:
        response = HttpResponse(content_type='text/csv')
        filename = f'reports_{datetime.datetime.now().strftime("%Y%m%d")}.csv'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        writer = csv.writer(response)
        writer.writerow(["Student", "Category", "Report Type", "Status", "Generated At", "Data"])
        for report in queryset:
            data_str = (report.data if isinstance(report.data, str)
                        else json.dumps(report.data, ensure_ascii=False))
            writer.writerow([
                report.student.full_name if report.student else "No Student",
                report.get_category_display(),
                report.report_type,
                report.status,
                report.generated_at,
                data_str
            ])
        return response
    except Exception as e:
        logger.exception("Error exporting reports to CSV: %s", e)
        return HttpResponse("An error occurred while generating the CSV file.", status=500)

def export_reports_word(modeladmin, request, queryset):
    """
    Admin action to export selected reports as a Word document.
    """
    try:
        doc = Document()
        doc.add_heading("Report Information", level=1)
        for index, report in enumerate(queryset):
            doc.add_paragraph(f"Report for {report.student.full_name} - {report.get_category_display()}")
            if index < len(queryset) - 1:
                doc.add_page_break()
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = 'attachment; filename="reports.docx"'
        doc.save(response)
        return response
    except Exception as e:
        logger.exception("Error exporting reports to Word: %s", e)
        return HttpResponse("An error occurred while generating the Word document.", status=500)

class ReportAdmin(admin.ModelAdmin):
    list_display = ('student', 'get_category_display', 'report_type', 'status', 'generated_at')
    list_filter = ('category', 'status')
    search_fields = ('student__full_name', 'report_type', 'category')
    actions = [export_reports_pdf, export_reports_csv, export_reports_word]
    readonly_fields = ('student', 'category', 'report_type', 'status', 'generated_at', 'pretty_data')
    fields = ('student', 'category', 'report_type', 'status', 'generated_at', 'pretty_data')

    def pretty_data(self, obj):
        """Display JSON data in a formatted HTML table."""
        try:
            if isinstance(obj.data, dict):
                data_dict = obj.data
            else:
                data_dict = json.loads(obj.data)
        except Exception:
            return obj.data

        html = "<table style='max-width:600px; border-collapse: collapse;'>"
        for key, value in data_dict.items():
            if value is None:
                value = "<em>Not Available</em>"
            html += f"<tr style='border:1px solid #ccc;'><th style='padding:4px; text-align:left; background:#f5f5f5;'>{key}</th><td style='padding:4px;'>{value}</td></tr>"
        html += "</table>"
        return mark_safe(html)
    pretty_data.short_description = "Report Data"

    def has_add_permission(self, request):
        # Prevent manual addition via admin.
        return False

admin.site.register(Report, ReportAdmin)

class ReportsDashboardAdmin(admin.ModelAdmin):
    """
    Dummy admin for ReportsDashboard. Redirects to the reports dashboard.
    """
    def changelist_view(self, request, extra_context=None):
        return redirect("/reports/")

    def has_add_permission(self, request):
        return False

    def has_delete_permission(self, request, obj=None):
        return False

admin.site.register(ReportsDashboard, ReportsDashboardAdmin)
