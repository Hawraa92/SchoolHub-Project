"""
Django Admin settings for SchoolHub's student management system.

This file configures export actions (PDF, Excel, Word, CSV) and ModelAdmin settings
for various student-related models.
"""

import datetime
import logging
import csv
from io import BytesIO

import pandas as pd
from django.contrib import admin
from django.http import HttpResponse
from django.urls import reverse
from django.utils.html import format_html
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from .models import (
    Student,
    ChronicIllness,
    Subject,
    Grade,
    EconomicSituation,
    SocialMediaAndTechnology,
    HealthInformation,
    GradeHistory,
    StudentPerformanceTrend,
)
from teachers.models import Teacher

logger = logging.getLogger(__name__)


# =============================================================================
# Export Functions
# =============================================================================
def export_student_pdf(modeladmin, request, queryset):
    """
    Export selected students as a PDF document.
    """
    try:
        response = HttpResponse(content_type='application/pdf')
        filename = f'students_{datetime.datetime.now().strftime("%Y%m%d")}.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        pdf = canvas.Canvas(response, pagesize=letter)
        pdf.setTitle("Student Information")
        y = 750
        for student in queryset:
            # New page if nearing the bottom
            if y < 100:
                pdf.showPage()
                y = 750
            pdf.drawString(100, y, f"Student Name: {student.full_name}")
            y -= 20
        pdf.save()
        return response
    except Exception as e:
        logger.exception("Error exporting students to PDF: %s", e)
        return HttpResponse("An error occurred while generating the PDF.", status=500)


def export_students_excel(modeladmin, request, queryset):
    """
    Export selected students as an Excel file.
    """
    try:
        data = [
            {
                "Full Name": student.full_name,
                "Email": student.email,
                "Mobile": student.mobile,
                "Date of Birth": student.date_of_birth,
                "Grade Level": student.grade_level or "None",
                "Academic Performance": student.academic_performance or "None",
            }
            for student in queryset
        ]
        df = pd.DataFrame(data)
        with BytesIO() as buffer:
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name="Students")
            buffer.seek(0)
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            response['Content-Disposition'] = 'attachment; filename="students.xlsx"'
            return response
    except Exception as e:
        logger.exception("Error exporting students to Excel: %s", e)
        return HttpResponse("An error occurred while generating the Excel file.", status=500)


def export_students_word(modeladmin, request, queryset):
    """
    Export selected students as a Word document.
    """
    try:
        doc = Document()
        doc.add_heading("Student Information", level=1)
        for index, student in enumerate(queryset):
            doc.add_paragraph(f"Student Name: {student.full_name}")
            # Add a page break if there are more than one student
            if index < len(queryset) - 1:
                doc.add_page_break()
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = 'attachment; filename="students.docx"'
        doc.save(response)
        return response
    except Exception as e:
        logger.exception("Error exporting students to Word: %s", e)
        return HttpResponse("An error occurred while generating the Word document.", status=500)


def export_students_csv(modeladmin, request, queryset):
    """
    Export selected students as a CSV file with selected fields for analysis.
    The CSV includes academic performance, attendance, and additional indicators.
    """
    try:
        response = HttpResponse(content_type='text/csv')
        filename = f'students_{datetime.datetime.now().strftime("%Y%m%d")}.csv'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        writer = csv.writer(response)
        
        writer.writerow([
            "Full Name",
            "Academic Performance",
            "Attendance Percentage",
            "Seat Zone",
            "Grades",
            "Academic Stress",           # from HealthInformation
            "Motivation",                # from HealthInformation
            "Depression",                # from HealthInformation
            "Sleep Disorder",            # from HealthInformation
            "Study Life Balance",        # from HealthInformation
            "Family Pressures",          # from HealthInformation
            "Parents Marital Status",    # from EconomicSituation
            "Family Income Level",       # from EconomicSituation
            "Housing Status",            # from EconomicSituation
            "Has Private Study Room",    # from EconomicSituation
            "Daily Food Availability",   # from EconomicSituation
            "Has School Uniform",        # from EconomicSituation
            "Has Stationery",            # from EconomicSituation
            "Receives Private Tutoring", # from EconomicSituation
            "Daily Study Hours",         # from EconomicSituation
            "Works After School",        # from EconomicSituation
            "Has Electronic Device",     # from SocialMediaAndTechnology
            "Device Usage Purpose",      # from SocialMediaAndTechnology
            "Has Social Media Accounts", # from SocialMediaAndTechnology
            "Daily Screen Time",         # from SocialMediaAndTechnology
            "Social Media Impact On Studies",  # from SocialMediaAndTechnology
            "Content Type Watched",      # from SocialMediaAndTechnology
            "Plays Video Games",         # from SocialMediaAndTechnology
            "Daily Gaming Hours"         # from SocialMediaAndTechnology
        ])
        
        for student in queryset:
            # Attempt to get GPA (if available)
            try:
                gpa = student.studentperformancetrend.gpa if hasattr(student, 'studentperformancetrend') else "N/A"
            except Exception:
                gpa = "N/A"
            
            # Combine grades into a single string
            try:
                grades = "; ".join([f"{grade.subject.name}:{grade.score}" for grade in student.grades.all()])
            except Exception:
                grades = "N/A"
            
            health_info = getattr(student, 'health_information', None)
            academic_stress = health_info.academic_stress if health_info else "N/A"
            motivation = health_info.motivation if health_info else "N/A"
            depression = health_info.depression if health_info else "N/A"
            sleep_disorder = health_info.sleep_disorder if health_info else "N/A"
            study_life_balance = health_info.study_life_balance if health_info else "N/A"
            family_pressures = health_info.family_pressures if health_info else "N/A"
            
            econ_info = getattr(student, 'economic_situation', None)
            parents_marital_status = econ_info.parents_marital_status if econ_info else "N/A"
            family_income_level = econ_info.family_income_level if econ_info else "N/A"
            housing_status = econ_info.housing_status if econ_info else "N/A"
            has_private_study_room = econ_info.has_private_study_room if econ_info else "N/A"
            daily_food_availability = econ_info.daily_food_availability if econ_info else "N/A"
            has_school_uniform = econ_info.has_school_uniform if econ_info else "N/A"
            has_stationery = econ_info.has_stationery if econ_info else "N/A"
            receives_private_tutoring = econ_info.receives_private_tutoring if econ_info else "N/A"
            daily_study_hours = econ_info.daily_study_hours if econ_info else "N/A"
            works_after_school = econ_info.works_after_school if econ_info else "N/A"
            
            tech_info = getattr(student, 'tech_and_social', None)
            has_electronic_device = tech_info.has_electronic_device if tech_info else "N/A"
            device_usage_purpose = tech_info.device_usage_purpose if tech_info else "N/A"
            has_social_media_accounts = tech_info.has_social_media_accounts if tech_info else "N/A"
            daily_screen_time = tech_info.daily_screen_time if tech_info else "N/A"
            social_media_impact_on_studies = tech_info.social_media_impact_on_studies if tech_info else "N/A"
            content_type_watched = tech_info.content_type_watched if tech_info else "N/A"
            plays_video_games = tech_info.plays_video_games if tech_info else "N/A"
            daily_gaming_hours = tech_info.daily_gaming_hours if tech_info else "N/A"
            
            writer.writerow([
                student.full_name,
                student.academic_performance,
                student.attendance_percentage,
                student.seat_zone,
                grades,
                academic_stress,
                motivation,
                depression,
                sleep_disorder,
                study_life_balance,
                family_pressures,
                parents_marital_status,
                family_income_level,
                housing_status,
                has_private_study_room,
                daily_food_availability,
                has_school_uniform,
                has_stationery,
                receives_private_tutoring,
                daily_study_hours,
                works_after_school,
                has_electronic_device,
                device_usage_purpose,
                has_social_media_accounts,
                daily_screen_time,
                social_media_impact_on_studies,
                content_type_watched,
                plays_video_games,
                daily_gaming_hours,
            ])
        return response
    except Exception as e:
        logger.exception("Error exporting students to CSV: %s", e)
        return HttpResponse("An error occurred while generating the CSV file.", status=500)


# =============================================================================
# Inline Admins
# =============================================================================
class GradeInline(admin.TabularInline):
    model = Grade
    extra = 1
    fields = (
        'subject',
        'score',
        'max_score',
        'percentage',
        'grade_level',
        'exam_type',
        'semester',
        'teacher',
        'date_recorded',
        'weight',
        'gpa_points',
    )
    readonly_fields = ('percentage', 'grade_level', 'gpa_points', 'date_recorded')
    show_change_link = True
    autocomplete_fields = ('subject',)


class HealthInformationInline(admin.StackedInline):
    model = HealthInformation
    can_delete = False
    verbose_name_plural = "Health Information"
    fk_name = 'student'
    fieldsets = (
        ('Physical Health', {
            'fields': (
                'has_chronic_illness',
                'general_health_status',
                'last_medical_checkup',
                'weight',
                'height',
            )
        }),
        ('Psychological/Behavioral', {
            'fields': (
                'academic_stress',
                'motivation',
                'depression',
                'sleep_disorder',
                'study_life_balance',
                'family_pressures',
            )
        }),
    )


class EconomicSituationInline(admin.StackedInline):
    model = EconomicSituation
    can_delete = False
    verbose_name_plural = "Economic Situation"
    fk_name = 'student'
    fieldsets = (
        ("Economic Information", {
            "fields": (
                'father_occupation',
                'mother_occupation',
                'parents_marital_status',
                'family_income_level',
                'income_source',
                'monthly_expenses',
                'housing_status',
                'access_to_electricity',
                'has_access_to_water',
                'access_to_internet',
                'has_private_study_room',
                'number_of_rooms_in_home',
                'daily_food_availability',
                'has_school_uniform',
                'has_stationery',
                'receives_scholarship',
                'receives_private_tutoring',
                'daily_study_hours',
                'works_after_school',
                'work_hours_per_week',
                'responsible_for_household_tasks',
                'transportation_mode',
                'distance_to_school',
                'has_health_insurance',
                'household_size',
                'sibling_rank',
            )
        }),
    )


class SocialMediaAndTechnologyInline(admin.StackedInline):
    model = SocialMediaAndTechnology
    can_delete = False
    verbose_name_plural = "Social Media and Technology"
    fk_name = 'student'


# =============================================================================
# Model Admins
# =============================================================================
@admin.register(Student)
class StudentAdmin(admin.ModelAdmin):
    list_display = (
        'full_name',
        'student_id',
        'grade_level',
        'email',
        'mobile',
        'academic_performance',
        'get_age',
        'generate_report_link',
    )
    list_display_links = ('full_name',)
    list_filter = ('grade_level', 'academic_performance', 'gender', 'nationality',)
    search_fields = (
        'full_name',
        'email',
        'mobile',
        'grades__subject__name',
        'grades__teacher__full_name',
        'grades__history__reason_for_update',
    )
    fieldsets = (
        ('Personal Information', {
            'fields': (
                'user',
                'full_name',
                'date_of_birth',
                'student_id',
                'get_age',
                'gender',
                'nationality',
                'address',
                'profile_image',
                'email',
                'mobile',
                'emergency_contact_name',
                'emergency_contact',
                'enrollment_date',
            )
        }),
        ('Guardian Information', {
            'fields': (
                'guardian_relationship',
                'guardian_address',
                'guardian_employment_status',
                'guardian_monthly_income',
                'guardian_education',
            )
        }),
        ('Academic Information', {
            'fields': (
                'grade_level',
                'subjects',
                'attendance_percentage',
                'academic_performance',
                'awards',
                'seat_zone',
            )
        }),
    )
    readonly_fields = ('get_age', 'student_id', 'academic_performance')
    filter_horizontal = ('subjects',)
    ordering = ('student_id',)
    inlines = [
        GradeInline,
        HealthInformationInline,
        EconomicSituationInline,
        SocialMediaAndTechnologyInline,
    ]
    list_select_related = ('user',)
    actions = [export_students_csv, export_student_pdf, export_students_excel, export_students_word]

    @admin.display(description="Age")
    def get_age(self, obj):
        return obj.age or "-"

    def generate_report_link(self, obj):
        """
        Returns a link to generate a report for the student using the new reports app.
        """
        if obj.pk:
            url = reverse("generate_single_student_report", args=[obj.pk])
            return format_html('<a class="button" href="{}" target="_blank">Generate Report</a>', url)
        return "-"
    generate_report_link.short_description = "Generate Report"


@admin.register(EconomicSituation)
class EconomicSituationAdmin(admin.ModelAdmin):
    pass


@admin.register(SocialMediaAndTechnology)
class SocialMediaAndTechnologyAdmin(admin.ModelAdmin):
    pass


@admin.register(ChronicIllness)
class ChronicIllnessAdmin(admin.ModelAdmin):
    list_display = ('name',)
    search_fields = ('name',)


@admin.register(Subject)
class SubjectAdmin(admin.ModelAdmin):
    list_display = ('name', 'description', 'credit_hours')
    search_fields = ('name',)


@admin.register(Grade)
class GradeAdmin(admin.ModelAdmin):
    list_display = (
        'student',
        'subject',
        'score',
        'max_score',
        'percentage',
        'get_grade_level',
        'exam_type',
        'get_semester',
        'teacher',
        'date_recorded',
        'weight',
        'gpa_points',
    )
    list_filter = ('exam_type',)
    search_fields = ('student__full_name', 'subject__name', 'teacher__full_name')
    ordering = ('-date_recorded',)
    autocomplete_fields = ('student', 'subject')
    list_select_related = ('student', 'subject', 'teacher')

    @admin.display(description="Grade Level")
    def get_grade_level(self, obj):
        return getattr(obj, 'grade_level', "-")

    @admin.display(description="Semester")
    def get_semester(self, obj):
        return getattr(obj, 'semester', "-")


@admin.register(GradeHistory)
class GradeHistoryAdmin(admin.ModelAdmin):
    list_display = ('grade', 'updated_by', 'previous_score', 'new_score', 'updated_at')
    search_fields = (
        'grade__student__full_name',
        'grade__subject__name',
        'updated_by__full_name',
        'reason_for_update'
    )


@admin.register(StudentPerformanceTrend)
class StudentPerformanceTrendAdmin(admin.ModelAdmin):
    list_display = ('student', 'semester', 'average_percentage', 'gpa')
    search_fields = ('student__full_name',)
